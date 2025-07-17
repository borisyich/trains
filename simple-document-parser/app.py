import io
import logging
import nest_asyncio
import ollama
import openpyxl
import os
import pandas as pd
import PyPDF2
import streamlit as st
import sys
import time
import warnings
from docx import Document
from dotenv import load_dotenv
from pdf2image import convert_from_bytes
from PIL import Image
import pytesseract
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.docstore.document import Document as LangChainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from pathlib import Path

# Suppress warnings to keep logs clean and avoid cluttering the console
warnings.filterwarnings("ignore")

# Load environment variables from .env file for configuration
load_dotenv()
# Path to Tesseract executable for OCR processing
TESSERACT_PATH = Path(os.getenv("TESSERACT_PATH"))
pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
# Chunk size and overlap for text splitting in RAG pipeline
CHUNK_SIZE = int(os.getenv('CHUNK_SIZE'))
CHUNK_OVERLAP = int(os.getenv('CHUNK_OVERLAP'))
# Number of relevant chunks to retrieve for answering questions
N_CHUNKS = int(os.getenv('N_CHUNKS'))

# Enable nested asyncio event loops for compatibility with Streamlit
nest_asyncio.apply()

# Configure logging to file and console for debugging and monitoring
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log', encoding='utf-8'),
        logging.StreamHandler()  # Output logs to console for debugging
    ]
)
logger = logging.getLogger(__name__)

# Global exception handler to log uncaught exceptions
def handle_exception(exc_type, exc_value, exc_traceback):
    """Logs unhandled exceptions to the app.log file and console."""
    logger.error("Необработанное исключение:", exc_info=(exc_type, exc_value, exc_traceback))
    with open("app.log", "a", encoding="utf-8") as log_file:
        log_file.write(f"Необработанное исключение: {exc_value}\n")

sys.excepthook = handle_exception

def extract_text(file):
    """
    Extracts text from uploaded files of various formats (.txt, .pdf, .docx, .xlsx).
    """
    start_time = time.time()
    file_extension = os.path.splitext(file.name)[1].lower()
    logger.info(f"Начало извлечения текста из файла: {file.name} (тип: {file_extension})")
    text = ""
    
    try:
        if file_extension == ".txt":
            # Read text directly from .txt files
            text = file.read().decode("utf-8")
            logger.info(f"Текст успешно извлечен из .txt файла: {len(text)} символов")
        elif file_extension == ".pdf":
            # Attempt to extract text from PDF using PyPDF2
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            logger.info(f"Извлечено {len(text)} символов с помощью PyPDF2 из .pdf")
            
            # If text is empty or too short, fall back to OCR
            if not text.strip() or len(text.strip()) < 100:
                logger.info("Текст пустой или слишком короткий, применяю OCR")
                text = ""
                file.seek(0)  # Reset file pointer to start
                images = convert_from_bytes(file.read())
                for i, img in enumerate(images):
                    page_text = pytesseract.image_to_string(img, lang="eng+rus")
                    text += page_text + "\n"
                    logger.info(f"OCR: Извлечено {len(page_text)} символов со страницы {i+1}")
        elif file_extension == ".docx":
            # Extract text from .docx files using python-docx
            doc = Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
            logger.info(f"Текст успешно извлечен из .docx файла: {len(text)} символов")
        elif file_extension == ".xlsx":
            # Extract text from .xlsx files by iterating through cells
            workbook = openpyxl.load_workbook(file)
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    text += " ".join([str(cell) for cell in row if cell is not None]) + "\n"
            logger.info(f"Текст успешно извлечен из .xlsx файла: {len(text)} символов")
        else:
            # Handle unsupported file formats
            logger.error(f"Неподдерживаемый формат файла: {file_extension}")
            st.error("Поддерживаются только файлы .txt, .pdf, .docx и .xlsx")
            return None
        logger.info(f"Извлечение текста завершено за {time.time() - start_time:.2f} секунд")
        return text.strip() if text.strip() else None
    except Exception as e:
        logger.error(f"Ошибка при обработке файла {file.name}: {str(e)}")
        st.error(f"Ошибка при обработке файла: {str(e)}")
        return None

def query_model(context, question):
    """
    Processes a user question using a Retrieval-Augmented Generation (RAG) pipeline.
    """
    start_time = time.time()
    logger.info(f"Отправка запроса к модели с вопросом: {question[:50]}...")
    
    try:
        embeddings = st.session_state.embeddings
        
        # Create or update vector store if document text has changed
        if st.session_state.vector_store is None or st.session_state.document_text != context:
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=CHUNK_SIZE,
                chunk_overlap=CHUNK_OVERLAP,
                length_function=len
            )
            chunks = text_splitter.split_text(context)
            documents = [LangChainDocument(page_content=chunk) for chunk in chunks]
            logger.info(f"Текст разбит на {len(chunks)} чанков")
            st.session_state.vector_store = Chroma.from_documents(documents, embeddings)
            logger.info("Векторное хранилище создано")

        # Retrieve relevant document chunks for the question
        retriever = st.session_state.vector_store.as_retriever(search_kwargs={"k": N_CHUNKS})
        relevant_docs = retriever.get_relevant_documents(question)
        context_for_prompt = "\n".join([doc.page_content for doc in relevant_docs])
        logger.info(f"Найдено {len(relevant_docs)} релевантных чанков")

        # Construct prompt for the language model
        prompt = (
            "Ты — полезный ассистент, который отвечает на вопросы на основе предоставленного документа. "
            "Отвечай кратко и по делу, используя только релевантную информацию из контекста.\n\n"
            f"Контекст документа:\n{context_for_prompt}\n\n"
            f"Вопрос: {question}\nОтвет:"
        )

        # Query the Gemma3n model via Ollama
        response = ollama.chat(
            model="gemma3n",
            messages=[
                {"role": "system", "content": "Ты — полезный ассистент, который отвечает на вопросы на основе предоставленного документа. Отвечай кратко и по делу."},
                {"role": "user", "content": prompt}
            ]
        )
        answer = response['message']['content'].strip()
        logger.info(f"Получен ответ: {len(answer)} символов")
        
        logger.info(f"Запрос к модели завершен за {time.time() - start_time:.2f} секунд")
        return answer
    except Exception as e:
        logger.error(f"Ошибка при обработке запроса: {str(e)}")
        return f"[Ошибка при обработке запроса: {str(e)}]"

# Initialize session state variables for document text, history, embeddings, and vector store
if "document_text" not in st.session_state:
    st.session_state.document_text = None
if "history" not in st.session_state:
    st.session_state.history = []
if "embeddings" not in st.session_state:
    # Initialize HuggingFace embeddings for text vectorization
    st.session_state.embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
        model_kwargs={'device': 'cpu'}
    )
    logger.info("Модель эмбеддингов инициализирована")
if "vector_store" not in st.session_state:
    st.session_state.vector_store = None

# Streamlit UI setup
st.title("Simple Documents Parser")
st.write("Загрузите документ (.txt, .pdf, .docx или .xlsx) и задавайте вопросы по его содержимому.")

# File uploader for document input
uploaded_file = st.file_uploader("Выберите документ", type=["txt", "pdf", "docx", "xlsx"])

# Process uploaded file and extract text
if uploaded_file is not None and st.session_state.document_text is None:
    logger.info(f"Файл загружен: {uploaded_file.name}")
    document_text = extract_text(uploaded_file)
    if document_text:
        st.session_state.document_text = document_text
        st.success(f"Документ успешно загружен! Длина текста: {len(document_text)} символов")
        logger.info(f"Документ загружен, длина текста: {len(document_text)} символов")
    else:
        st.session_state.document_text = None
        logger.warning("Не удалось извлечь текст из файла")

# Input field for user questions
question = st.text_input("Задайте вопрос по документу:")

# Button to submit question and get response
if st.button("Отправить вопрос", key="submit_question"):
    if st.session_state.document_text and question:
        with st.spinner("Обрабатываю запрос..."):
            logger.info("Нажата кнопка 'Отправить вопрос'")
            answer = query_model(st.session_state.document_text, question)
            # Append question and answer to history
            st.session_state.history.append({"Вопрос": question, "Ответ": answer})
            st.write("**Ответ:**")
            st.write(answer)
            logger.info("Ответ успешно отображен в интерфейсе")
    else:
        st.error("Пожалуйста, загрузите документ и введите вопрос.")
        logger.warning("Ошибка: отсутствует документ или вопрос")

# Display question and answer history
if st.session_state.history:
    st.markdown("---")
    st.subheader("История вопросов и ответов")
    history_df = pd.DataFrame(st.session_state.history)
    st.table(history_df)
    logger.info(f"Отображена история: {len(st.session_state.history)} записей")

# Button to clear question history
if st.button("Очистить историю", key="clear_history"):
    st.session_state.history = []
    st.success("История очищена!")
    logger.info("История вопросов и ответов очищена")

# Instructions for using the application
st.markdown("---")
st.markdown("**Инструкция:**\n1. Загрузите документ в формате .txt, .pdf, .docx или .xlsx.\n2. Введите вопрос, связанный с содержимым документа.\n3. Нажмите 'Отправить вопрос', чтобы получить ответ от модели Gemma3n.\n4. Просмотрите историю вопросов и ответов ниже.\n5. Нажмите 'Очистить историю', чтобы удалить все записи.")

# Display recent application logs
st.markdown("---")
st.subheader("Логи приложения")
if os.path.exists("app.log"):
    with open("app.log", "r", encoding="utf-8") as log_file:
        logs = log_file.readlines()
        # Show last 10 lines of logs
        st.text_area("Последние логи", value="".join(logs[-10:]), height=200, disabled=True)
    logger.info("Логи отображены в интерфейсе")
else:
    st.text("Логи пока недоступны")
    logger.info("Файл логов еще не создан")